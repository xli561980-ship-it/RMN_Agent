#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ingest.py — 实验室文献与操作手册：解析、分层切分、LLM 元数据增强、向量化入库

流程概览：
1. **目录分流**：`data/papers/` 下为论文（Paper），`data/manuals/` 下为操作手册（SOP）；二者分别入库。
2. 支持 **PDF** 与 **Word（.docx）**：LlamaParse 为主（保留标题层级）；PDF 可选 pdfplumber 兜底；Word 可选 python-docx 按样式转 Markdown。
3. 论文目录内：正文（Main）与补充材料（SI）自动配对；手册目录内不做 SI 配对。
4. **差异化切分**：论文 `doc_type=paper`，第二层 chunk 800/150；手册 `doc_type=sop`，第二层 chunk 1200/200，以保留完整步骤。
5. 分层切分：MarkdownHeaderTextSplitter → RecursiveCharacterTextSplitter（chunk 参数按类型不同）。
6. 论文 chunk 含 project_id、doc_role、doc_type=paper 等；手册 chunk 含 doc_type=sop、doc_role=manual 等。
   仅论文 chunk 在 `page_content` 首行写入简短 `[DOC] title: … | source: … | project_id: … | doc_role: …` 前缀，使题录参与向量匹配；**需对已入库文件增量/全量重跑 ingest**（或 `REBUILD_CHROMA=1`）后嵌入向量才含此前缀。
7. **source 元数据**：使用相对 `data/` 的路径（如 `papers/a.pdf`、`manuals/SOP.docx`）作为唯一键，支持增量删除与去重。

环境变量（补充）：
- GOOGLE_API_KEY / GEMINI_API_KEY：用于 LLM 元数据抽取（与嵌入可共用）。
- GOOGLE_LLM_MODEL：元数据抽取所用对话模型，默认 gemini-3-flash-preview（可与 GOOGLE_EMBEDDING_MODEL 不同）。
- GOOGLE_EMBEDDING_MODEL：Google 向量模型，默认 `gemini-embedding-001`（专用嵌入模型，勿用对话模型名）。
- INGEST_METADATA_EXCERPT_CHARS：送入 LLM 的最大字符数，默认 14000。
- 增量更新：`processed_files.json` 以「相对 data 的路径」为键记录 MD5/mtime；
  `REBUILD_CHROMA=1` 时清空向量库并重置记录。
- Word：本地解析依赖 `python-docx`；长文档可按 `DOCX_SEGMENT_CHARS`（默认 8000）切段并赋虚拟页码。
  旧版 `.doc` 未支持，请在 Word 中另存为 `.docx`。
"""

from __future__ import annotations

import hashlib
import json
import math
import os
import re
import shutil
import sys
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional, Sequence, Tuple

from dotenv import load_dotenv

ROOT = Path(__file__).resolve().parent
load_dotenv(ROOT / ".env")

from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings

try:
    from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
except ImportError:  # pragma: no cover
    ChatGoogleGenerativeAI = None  # type: ignore
    GoogleGenerativeAIEmbeddings = None  # type: ignore

try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
except ImportError:  # pragma: no cover
    HuggingFaceEmbeddings = None  # type: ignore

import pdfplumber

try:
    from llama_parse import LlamaParse
except ImportError:  # pragma: no cover
    LlamaParse = None  # type: ignore

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None  # type: ignore

from pydantic import BaseModel, ConfigDict, Field

from chunking import ChunkingConfig, split_document


# =============================================================================
# 路径与超参数
# =============================================================================
DATA_DIR = Path(__file__).resolve().parent / "data"
DATA_PAPERS_DIR = DATA_DIR / "papers"
DATA_MANUALS_DIR = DATA_DIR / "manuals"
CHROMA_PERSIST_DIR = Path(os.getenv("CHROMA_PERSIST_DIR", "./chroma_db")).resolve()
CHROMA_COLLECTION_NAME = os.getenv("CHROMA_COLLECTION_NAME", "lab_literature_rag")
# 已成功向量化的文档指纹（键 = 相对 data/ 的路径，与 chunk metadata["source"] 对齐）
PROCESSED_FILES_JSON = Path(__file__).resolve().parent / "processed_files.json"
CORPUS_MANIFEST_JSON = Path(__file__).resolve().parent / "corpus_manifest.json"

# 当前支持的文档扩展名（小写，含点）
INGEST_SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx"})

# 论文 / SOP：可配置 chunking，默认保留原有 header-aware 思路并让 SOP 走 parent-child。
CHUNK_STRATEGY = os.getenv("CHUNK_STRATEGY", "header_aware")
PAPER_CHUNK_STRATEGY = os.getenv("PAPER_CHUNK_STRATEGY", CHUNK_STRATEGY or "header_aware")
SOP_CHUNK_STRATEGY = os.getenv("SOP_CHUNK_STRATEGY", CHUNK_STRATEGY or "parent_child")
PAPER_CHUNK_SIZE = int(os.getenv("PAPER_CHUNK_SIZE", "900"))
PAPER_CHUNK_OVERLAP = int(os.getenv("PAPER_CHUNK_OVERLAP", "150"))
SOP_CHILD_CHUNK_SIZE = int(os.getenv("SOP_CHILD_CHUNK_SIZE", os.getenv("SOP_CHUNK_SIZE", "500")))
SOP_CHILD_CHUNK_OVERLAP = int(os.getenv("SOP_CHILD_CHUNK_OVERLAP", os.getenv("SOP_CHUNK_OVERLAP", "80")))
SOP_PARENT_CHUNK_SIZE = int(os.getenv("SOP_PARENT_CHUNK_SIZE", "1800"))
SOP_PARENT_CHUNK_OVERLAP = int(os.getenv("SOP_PARENT_CHUNK_OVERLAP", "200"))
SOP_CHUNK_SIZE = SOP_CHILD_CHUNK_SIZE
SOP_CHUNK_OVERLAP = SOP_CHILD_CHUNK_OVERLAP

INGEST_METADATA_EXCERPT_CHARS = int(os.getenv("INGEST_METADATA_EXCERPT_CHARS", "14000"))

HEADERS_TO_SPLIT_ON: List[tuple[str, str]] = [
    ("#", "Header 1"),
    ("##", "Header 2"),
    ("###", "Header 3"),
]


class LocalHashEmbeddings:
    """Small deterministic embedding fallback for offline tests and smoke benchmarks."""

    def __init__(self, dimensions: int = 384) -> None:
        self.dimensions = max(16, int(dimensions))

    def _embed(self, text: str) -> List[float]:
        vec = [0.0] * self.dimensions
        tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9+\-_/]{1,}|[\u4e00-\u9fff]{2,}", text or "")
        for token in tokens:
            digest = hashlib.md5(token.casefold().encode("utf-8")).digest()
            idx = int.from_bytes(digest[:4], "little") % self.dimensions
            sign = 1.0 if digest[4] % 2 == 0 else -1.0
            vec[idx] += sign
        norm = math.sqrt(sum(x * x for x in vec)) or 1.0
        return [x / norm for x in vec]

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [self._embed(t) for t in texts]

    def embed_query(self, text: str) -> List[float]:
        return self._embed(text)


def pdf_source_key(file_path: Path) -> str:
    """
    Chroma 中 metadata「source」唯一键：相对于 data/ 的 POSIX 路径（如 papers/a.pdf、manuals/SOP.docx）。
    避免 papers/ 与 manuals/ 下同名文件冲突，并与增量删除逻辑一致。
    """
    root = DATA_DIR.resolve()
    try:
        return file_path.resolve().relative_to(root).as_posix()
    except ValueError:
        return file_path.name


def is_supported_ingest_file(path: Path) -> bool:
    """是否为本流水线支持的论文/手册文件类型。"""
    return path.is_file() and path.suffix.lower() in INGEST_SUPPORTED_SUFFIXES


def ensure_ingest_directories() -> None:
    """确保 data、papers、manuals 子目录存在。"""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    DATA_PAPERS_DIR.mkdir(parents=True, exist_ok=True)
    DATA_MANUALS_DIR.mkdir(parents=True, exist_ok=True)


def iter_ingest_jobs() -> List[Tuple[Path, Literal["paper", "sop"]]]:
    """
    枚举待入库文档（.pdf / .docx）及其类型（论文 / 手册）。
    仅扫描 `data/papers` 与 `data/manuals`。
    """
    jobs: List[Tuple[Path, Literal["paper", "sop"]]] = []
    if DATA_PAPERS_DIR.is_dir():
        for p in sorted(DATA_PAPERS_DIR.iterdir()):
            if is_supported_ingest_file(p):
                jobs.append((p, "paper"))
    if DATA_MANUALS_DIR.is_dir():
        for p in sorted(DATA_MANUALS_DIR.iterdir()):
            if is_supported_ingest_file(p):
                jobs.append((p, "sop"))
    return jobs


def iter_pdf_ingest_jobs() -> List[Tuple[Path, Literal["paper", "sop"]]]:
    """兼容旧名，等价于 iter_ingest_jobs。"""
    return iter_ingest_jobs()


def build_sop_global_metadata(pdf_path: Path) -> Dict[str, Any]:
    """手册：不调用论文题录 LLM，使用文件名回退 + 强制 doc_type / doc_role。"""
    meta_llm, pid = _fallback_metadata_from_path(pdf_path)
    flat = _flatten_paper_meta_for_chroma(meta_llm, pid)
    flat["doc_type"] = "sop"
    flat["doc_role"] = "manual"
    flat["paired_supplementary_sources"] = ""
    flat["paired_main_source"] = ""
    return flat


# ---------- 补充材料文件名识别与配对 ----------

# 文件名（stem）中若命中下列子串，则视为补充材料/SI
SI_STEM_SUBSTRINGS: Tuple[str, ...] = (
    "supplementary",
    "supplement",
    "appendix",
    "supporting information",
    "supporting-information",
)


def is_supplementary_filename(path: Path) -> bool:
    """
    判断文献文件是否为补充材料（SI）；适用于 .pdf / .docx 等（按 stem 判断）。

    规则（满足其一即可）：
    - stem 中含 supplementary / supplement / appendix / supporting information 等；
    - stem 在词边界上含 SI（如 `_SI`、`-SI`、`_si` 等），避免误伤如「design」类词。
    """
    stem_raw = path.stem
    s = stem_raw.lower()
    for sub in SI_STEM_SUBSTRINGS:
        if sub in s:
            return True
    if re.search(r"(^|[_\-\s])si([_\-\s]|$)", s, flags=re.IGNORECASE):
        return True
    if s.endswith(" si"):
        return True
    return False


def pairing_core_stem(path: Path) -> str:
    """
    用于 Main ↔ SI 配对的「规范化 stem」：去掉 SI/补充材料常见后缀，再转小写比较。

    例如：`Paper_ABC_SI` → `paper_abc`；`Paper Supplementary Information` → `paper`。
    """
    stem = path.stem
    stem = re.sub(
        r"(?i)[\s_-]*(supplementary|supplement)\s*(information|info|materials?)?[\s_-]*.*$",
        "",
        stem,
    )
    stem = re.sub(r"(?i)[\s_-]*appendix[\s_-]?[a-z0-9]*$", "", stem)
    stem = re.sub(r"(?i)[\s_-]*supporting[\s_-]*information[\s_-]*.*$", "", stem)
    stem = re.sub(r"(?i)(^|[_\-\s])si([\s_-]|$).*$", "", stem)
    stem = stem.strip(" _-.")
    return stem.lower() or path.stem.lower()


@dataclass
class PaperUnit:
    """一篇文献在磁盘上的组织单元：可选正文 + 若干补充材料文件。"""

    core_stem: str
    main_pdf: Optional[Path] = None
    extra_main_pdfs: List[Path] = field(default_factory=list)
    supplementary_pdfs: List[Path] = field(default_factory=list)


def group_pdfs_into_paper_units(paper_files: Sequence[Path]) -> List[PaperUnit]:
    """
    将 papers/ 下文献文件划分为单元并完成 Main/SI 配对（支持 .pdf / .docx）。

    - 每个「非 SI」文件作为该 core_stem 的 main；
    - 每个 SI 文件按 pairing_core_stem 归入同一单元；若无对应 main，则形成「仅有 SI」的单元（孤儿 SI）。
    """
    mains = [p for p in paper_files if not is_supplementary_filename(p)]
    sis = [p for p in paper_files if is_supplementary_filename(p)]

    units_by_core: Dict[str, PaperUnit] = {}
    for m in mains:
        core = pairing_core_stem(m)
        if core not in units_by_core:
            units_by_core[core] = PaperUnit(core_stem=core, main_pdf=m, supplementary_pdfs=[])
        else:
            # 极少见：同一 core 多个正文，保留先扫描到的，其余仅并入列表由人工整理文件名
            u = units_by_core[core]
            if u.main_pdf is None:
                u.main_pdf = m
            else:
                u.extra_main_pdfs.append(m)

    for s in sis:
        core = pairing_core_stem(s)
        if core not in units_by_core:
            units_by_core[core] = PaperUnit(core_stem=core, main_pdf=None, supplementary_pdfs=[])
        units_by_core[core].supplementary_pdfs.append(s)

    return sorted(units_by_core.values(), key=lambda u: u.core_stem)


def warn_paper_unit_pairing_anomalies(units: Sequence[PaperUnit]) -> None:
    """Surface SI pairing risks instead of silently accepting filename heuristics."""
    for unit in units:
        if unit.supplementary_pdfs and unit.main_pdf is None:
            print(
                f"  [配对警告] 单元 [{unit.core_stem}] 只有补充材料、没有正文；将用第一个待处理文件作为锚点。",
                file=sys.stderr,
            )
        if unit.extra_main_pdfs:
            extras = " | ".join(p.name for p in unit.extra_main_pdfs[:5])
            more = f" (+{len(unit.extra_main_pdfs) - 5} more)" if len(unit.extra_main_pdfs) > 5 else ""
            print(
                f"  [配对警告] 单元 [{unit.core_stem}] 发现多个疑似正文；主锚点为 "
                f"{unit.main_pdf.name if unit.main_pdf else '(none)'}，额外正文也会按 main_text 入库：{extras}{more}",
                file=sys.stderr,
            )


# ---------- LLM 元数据（Pydantic + Gemini） ----------


class PaperMetadataLLM(BaseModel):
    """LLM 结构化输出：每篇论文一份，用于写入各 chunk 的 metadata 以支持过滤。"""

    model_config = ConfigDict(extra="ignore")

    paper_title: str = Field(description="论文完整标题")
    authors: List[str] = Field(default_factory=list, description="全部作者列表，顺序与论文一致")
    publication_year: Optional[int] = Field(default=None, description="发表年份，四位整数，未知则 null")
    experiment_keywords: List[str] = Field(
        default_factory=list,
        description="实验相关关键词，如材料、方法、仪器、表征手段等，5~15 个为宜",
    )


def _slug_project_id(title: str, anchor_path: Path) -> str:
    """
    project_id：在论文标题基础上生成 URL 友好 slug，并拼接文件名哈希后缀，避免不同文件撞车。
    """
    t = (title or "").strip() or anchor_path.stem
    t = re.sub(r"\s+", "_", t)
    t = re.sub(r"[^\w\u4e00-\u9fff\-_.]", "", t)
    slug = (t[:100] or "untitled").strip("_")
    h = hashlib.md5(anchor_path.name.encode("utf-8")).hexdigest()[:8]
    return f"{slug}_{h}"


def _flatten_paper_meta_for_chroma(meta: PaperMetadataLLM, project_id: str) -> Dict[str, Any]:
    """转为 Chroma 友好的标量 metadata（列表用分隔符拼成字符串，便于 where 过滤）。"""
    authors_s = " | ".join(a.strip() for a in meta.authors if (a or "").strip())
    kw_s = " | ".join(k.strip() for k in meta.experiment_keywords if (k or "").strip())
    out: Dict[str, Any] = {
        "project_id": project_id,
        "paper_title": (meta.paper_title or "").strip() or "unknown_title",
        "authors": authors_s[:4000],
        "experiment_keywords": kw_s[:4000],
    }
    if meta.publication_year is not None:
        out["publication_year"] = int(meta.publication_year)
    return out


def _build_metadata_llm():
    if ChatGoogleGenerativeAI is None:
        raise RuntimeError("未安装 langchain-google-genai，无法进行 LLM 元数据抽取。")
    api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("请配置 GOOGLE_API_KEY 或 GEMINI_API_KEY 以进行论文元数据抽取。")
    model = os.getenv("GOOGLE_LLM_MODEL", "gemini-3-flash-preview")
    return ChatGoogleGenerativeAI(model=model, google_api_key=api_key, temperature=0.0)


def excerpt_text_for_metadata(segments: Sequence[Document], max_chars: int) -> str:
    """将解析后的若干片段顺序拼接，截断到 max_chars，供 LLM 推断题录信息。"""
    parts: List[str] = []
    n = 0
    for seg in segments:
        piece = seg.page_content or ""
        if n + len(piece) > max_chars:
            parts.append(piece[: max_chars - n])
            break
        parts.append(piece)
        n += len(piece)
    return "\n\n".join(parts)


def extract_paper_metadata_llm(
    segments: Sequence[Document],
    anchor_path: Path,
    *,
    max_chars: int = INGEST_METADATA_EXCERPT_CHARS,
) -> Tuple[PaperMetadataLLM, str]:
    """
    调用 LLM 从正文摘录中抽取题录与实验关键词，并生成 project_id。

    若结构化输出失败，则回退为「文件名启发式」元数据，保证流水线不中断。
    """
    excerpt = excerpt_text_for_metadata(segments, max_chars)
    llm = _build_metadata_llm()
    structured = llm.with_structured_output(PaperMetadataLLM)

    sys_msg = (
        "你是学术文献元数据抽取助手。只根据用户给出的论文文本摘录推断题录信息；"
        "不得臆造不存在的作者或年份。若无法判断年份，publication_year 置为 null。"
        "experiment_keywords 用于后续实验类检索过滤，请优先从方法、材料、仪器、表征中抽取简短词或短语。"
    )
    human = f"文件名（仅供参考）: {anchor_path.name}\n\n论文摘录:\n{excerpt}"

    try:
        result = structured.invoke(
            [
                ("system", sys_msg),
                ("human", human),
            ]
        )
        if isinstance(result, PaperMetadataLLM):
            meta = result
        elif isinstance(result, dict):
            meta = PaperMetadataLLM(**result)  # type: ignore[misc]
        else:
            meta = PaperMetadataLLM.model_validate(result)  # type: ignore[attr-defined]
    except Exception as exc:  # pragma: no cover
        print(f"[警告] LLM 元数据抽取失败，使用文件名回退：{anchor_path.name} — {exc}", file=sys.stderr)
        meta = PaperMetadataLLM(
            paper_title=anchor_path.stem.replace("_", " "),
            authors=[],
            publication_year=None,
            experiment_keywords=[],
        )

    pid = _slug_project_id(meta.paper_title, anchor_path)
    return meta, pid


def _fallback_metadata_from_path(path: Path, project_id: Optional[str] = None) -> Tuple[PaperMetadataLLM, str]:
    """无正文、仅 SI 或 LLM 失败时的最弱回退。"""
    meta = PaperMetadataLLM(
        paper_title=path.stem.replace("_", " "),
        authors=[],
        publication_year=None,
        experiment_keywords=[],
    )
    pid = project_id or _slug_project_id(meta.paper_title, path)
    return meta, pid


# ---------- 解析与切分（与上一版一致，并增加 global_file_metadata） ----------


def _build_recursive_splitter(chunk_size: int, chunk_overlap: int) -> RecursiveCharacterTextSplitter:
    """第二层递归字符切分器；论文与手册使用不同 chunk_size / overlap。"""
    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", ". ", " ", ""],
        length_function=len,
        is_separator_regex=False,
    )


def _build_markdown_header_splitter() -> MarkdownHeaderTextSplitter:
    return MarkdownHeaderTextSplitter(
        headers_to_split_on=HEADERS_TO_SPLIT_ON,
        strip_headers=False,
        return_each_line=False,
    )


def _words_to_lines(words: List[dict]) -> str:
    if not words:
        return ""
    sorted_words = sorted(words, key=lambda w: (float(w["top"]), float(w["x0"])))
    lines: List[str] = []
    line_tokens: List[str] = []
    current_top: float | None = None
    line_tol = 2.5
    for w in sorted_words:
        top = float(w["top"])
        if current_top is None:
            current_top = top
        if abs(top - current_top) > line_tol and line_tokens:
            lines.append(" ".join(line_tokens))
            line_tokens = []
            current_top = top
        line_tokens.append(w.get("text") or "")
    if line_tokens:
        lines.append(" ".join(line_tokens))
    return "\n".join(lines)


def _extract_page_text_pdfplumber(page: pdfplumber.page.Page) -> str:
    words = page.extract_words(keep_blank_chars=False, use_text_flow=False)
    if not words:
        return (page.extract_text() or "").strip()
    width = float(page.width)
    mid = width / 2.0
    left_words: List[dict] = []
    right_words: List[dict] = []
    for w in words:
        x0, x1 = float(w["x0"]), float(w["x1"])
        xc = (x0 + x1) / 2.0
        (left_words if xc < mid else right_words).append(w)
    min_words_per_column = 15
    if len(left_words) >= min_words_per_column and len(right_words) >= min_words_per_column:
        left_txt = _words_to_lines(left_words).strip()
        right_txt = _words_to_lines(right_words).strip()
        if left_txt and right_txt:
            return f"{left_txt}\n\n{right_txt}"
        return left_txt or right_txt or ""
    fallback = (page.extract_text() or "").strip()
    joined = _words_to_lines(words).strip()
    return joined if len(joined) >= len(fallback) else fallback


def _pdfplumber_fallback_markdown_pages(pdf_path: Path, source_key: str) -> List[Document]:
    out: List[Document] = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            body = _extract_page_text_pdfplumber(page)
            if not body.strip():
                continue
            md = f"## 第 {i} 页\n\n{body.strip()}"
            out.append(Document(page_content=md, metadata={"source": source_key, "page": i}))
    return out


def parse_file_with_llama_parse(file_path: Path, source_key: str) -> List[Document]:
    """LlamaParse：支持 PDF、Word（.docx）等（以云端解析器实际支持为准），输出 Markdown 片段。"""
    if LlamaParse is None:
        raise RuntimeError("未安装 llama-parse。请执行: pip install llama-parse")
    api_key = os.getenv("LLAMA_CLOUD_API_KEY")
    if not api_key:
        raise RuntimeError("未设置 LLAMA_CLOUD_API_KEY，无法调用 LlamaParse。")

    parser = LlamaParse(api_key=api_key, result_type="markdown", verbose=False)
    raw_docs = parser.load_data(str(file_path))
    out: List[Document] = []
    for d in raw_docs:
        text = getattr(d, "text", None) or getattr(d, "get_content", lambda: "")()
        if not (text or "").strip():
            continue
        meta = dict(getattr(d, "metadata", None) or {})
        page_val = meta.get("page") or meta.get("page_label") or meta.get("page_number")
        try:
            page_int = int(page_val) if page_val is not None else 1
        except (TypeError, ValueError):
            page_int = 1
        out.append(
            Document(
                page_content=text.strip(),
                metadata={"source": source_key, "page": page_int},
            )
        )
    return out


def _load_pdf_markdown_segments(file_path: Path, source_key: str) -> List[Document]:
    """仅 PDF：LlamaParse 或 pdfplumber 兜底。"""
    allow_fb = (os.getenv("INGEST_PDFPLUMBER_FALLBACK", "") or "").lower() in ("1", "true", "yes")
    use_llama = bool(os.getenv("LLAMA_CLOUD_API_KEY")) and LlamaParse is not None

    if use_llama:
        try:
            return parse_file_with_llama_parse(file_path, source_key)
        except Exception as exc:
            if not allow_fb:
                raise
            print(f"[警告] LlamaParse 不可用或失败，已按兜底使用 pdfplumber：{file_path.name} — {exc}", file=sys.stderr)

    if not allow_fb:
        raise RuntimeError(
            "未配置 LLAMA_CLOUD_API_KEY 或未安装 llama-parse，且未开启 INGEST_PDFPLUMBER_FALLBACK。"
        )
    return _pdfplumber_fallback_markdown_pages(file_path, source_key)


def _docx_paragraphs_to_markdown(docx_path: Path) -> str:
    """
    使用 python-docx 将 .docx 转为近似 Markdown：按段落样式映射 Heading 1/2/3 为 # / ## / ###。
    """
    if DocxDocument is None:
        raise RuntimeError("未安装 python-docx，无法解析 .docx。请执行: pip install python-docx")
    doc = DocxDocument(str(docx_path))
    lines: List[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        style_name = (para.style.name if para.style else "") or ""
        st = style_name.strip().lower()
        if st in ("title",) or "heading 1" in st:
            lines.append("# " + t)
        elif "heading 2" in st:
            lines.append("## " + t)
        elif "heading 3" in st:
            lines.append("### " + t)
        else:
            lines.append(t)
    return "\n\n".join(lines).strip()


def _split_markdown_into_page_segments(body: str, source_key: str) -> List[Document]:
    """
    Word 无固定「页」概念：将长文按字符窗口切分为多个片段，page 从 1 递增，便于溯源展示。
    """
    max_seg = int(os.getenv("DOCX_SEGMENT_CHARS", "8000"))
    if len(body) <= max_seg:
        return [Document(page_content=body, metadata={"source": source_key, "page": 1})]
    out: List[Document] = []
    start = 0
    page = 1
    n = len(body)
    while start < n:
        end = min(start + max_seg, n)
        chunk_raw = body[start:end]
        if end < n:
            br = chunk_raw.rfind("\n\n")
            if br > max_seg // 4:
                chunk_raw = chunk_raw[: br + 2]
        chunk = chunk_raw.strip()
        if chunk:
            out.append(Document(page_content=chunk, metadata={"source": source_key, "page": page}))
            page += 1
        start += max(len(chunk_raw), 1)
    return out if out else [Document(page_content=body[:max_seg], metadata={"source": source_key, "page": 1})]


def _load_docx_markdown_segments(docx_path: Path, source_key: str) -> List[Document]:
    """Word：优先 LlamaParse；失败或未配置时回退 python-docx。"""
    use_llama = bool(os.getenv("LLAMA_CLOUD_API_KEY")) and LlamaParse is not None
    if use_llama:
        try:
            return parse_file_with_llama_parse(docx_path, source_key)
        except Exception as exc:
            print(f"[警告] LlamaParse 解析 Word 失败，回退本地 python-docx：{docx_path.name} — {exc}", file=sys.stderr)
    md = _docx_paragraphs_to_markdown(docx_path)
    if not md:
        raise RuntimeError(f"Word 文档为空或无可读文本：{docx_path.name}")
    return _split_markdown_into_page_segments(md, source_key)


def load_document_markdown_segments(file_path: Path, source_key: str) -> List[Document]:
    """统一入口：按扩展名选择 PDF / Word 解析链路。"""
    suf = file_path.suffix.lower()
    if suf == ".pdf":
        return _load_pdf_markdown_segments(file_path, source_key)
    if suf == ".docx":
        return _load_docx_markdown_segments(file_path, source_key)
    raise ValueError(f"不支持的文件类型: {file_path.suffix}（支持 {', '.join(sorted(INGEST_SUPPORTED_SUFFIXES))}）")


def load_pdf_markdown_segments(file_path: Path, source_key: str) -> List[Document]:
    """兼容旧函数名，等价于 load_document_markdown_segments。"""
    return load_document_markdown_segments(file_path, source_key)


def _clean_metadata(meta: Dict[str, Any]) -> Dict[str, Any]:
    cleaned: Dict[str, Any] = {}
    for k, v in meta.items():
        if v is None:
            continue
        if isinstance(v, (str, int, float, bool)):
            cleaned[k] = v
        else:
            cleaned[k] = str(v)
    return cleaned


def assess_text_quality(text: str) -> Dict[str, Any]:
    """
    Lightweight parse-quality heuristic for PDF/OCR chunks.

    Returns scalar values so they can be copied into Chroma metadata. This is
    not a perfect parser score; it simply flags common failure modes such as
    fused words from multi-column extraction or unreadably low whitespace.
    """
    body = text or ""
    chars = len(body)
    if chars == 0:
        return {"score": 0.0, "warning": "empty_chunk", "long_alpha_tokens": 0, "whitespace_ratio": 0.0}
    whitespace_ratio = sum(1 for ch in body if ch.isspace()) / chars
    long_alpha_tokens = len(re.findall(r"[A-Za-z]{35,}", body))
    replacement_chars = body.count("\ufffd")
    warnings: List[str] = []
    if chars > 300 and whitespace_ratio < 0.055:
        warnings.append("low_whitespace_ratio")
    if long_alpha_tokens >= 2:
        warnings.append("possible_fused_words")
    if replacement_chars:
        warnings.append("replacement_chars")
    penalty = 0.0
    penalty += min(0.45, long_alpha_tokens * 0.08)
    penalty += 0.25 if chars > 300 and whitespace_ratio < 0.055 else 0.0
    penalty += min(0.30, replacement_chars * 0.03)
    score = max(0.0, min(1.0, 1.0 - penalty))
    return {
        "score": round(score, 3),
        "warning": ",".join(warnings),
        "long_alpha_tokens": int(long_alpha_tokens),
        "whitespace_ratio": round(whitespace_ratio, 4),
    }


def _paper_embedding_prefix_line(meta: Dict[str, Any]) -> str:
    """
    Stable one-line prefix (English-heavy) prepended to each paper chunk body for embeddings.
    SOP / doc_type!=paper must not use this.
    """
    if meta.get("doc_type") != "paper":
        return ""
    raw_title = (meta.get("paper_title") or "").strip()
    if not raw_title or raw_title == "unknown_title":
        src_key = (meta.get("source") or "").strip()
        stem = Path(src_key).stem if src_key else ""
        raw_title = stem.replace("_", " ") if stem else "unknown_title"
    title_one = re.sub(r"[\r\n]+", " ", raw_title).strip()[:400]
    source_key = (meta.get("source") or "").strip() or "unknown_source"
    pid = (meta.get("project_id") or "").strip() or "unknown_project"
    role = (meta.get("doc_role") or "").strip() or "unknown_role"
    return f"[DOC] title: {title_one} | source: {source_key} | project_id: {pid} | doc_role: {role}\n"


def hierarchical_chunk_markdown_segment(
    segment: Document,
    header_splitter: MarkdownHeaderTextSplitter,
    recursive_splitter: RecursiveCharacterTextSplitter,
    *,
    global_file_metadata: Optional[Dict[str, Any]] = None,
    chunking_config: Optional[ChunkingConfig] = None,
) -> List[Document]:
    """
    双层切分：Markdown 标题 → 递归字数。

    global_file_metadata：与具体页/章节无关、但对每个 chunk 都必须相同的字段
    （project_id、doc_role、paper_title、authors、publication_year、experiment_keywords 等），
    在合并 metadata 时置于 Header 之后写入，确保不被章节 metadata 覆盖。
    """
    gfm = _clean_metadata(global_file_metadata or {})
    base = {
        **gfm,
        "source": segment.metadata.get("source", ""),
        "page": segment.metadata.get("page", 1),
    }
    if chunking_config is None:
        chunking_config = ChunkingConfig(
            strategy="header_aware",
            chunk_size=int(getattr(recursive_splitter, "_chunk_size", PAPER_CHUNK_SIZE)),
            chunk_overlap=int(getattr(recursive_splitter, "_chunk_overlap", PAPER_CHUNK_OVERLAP)),
        )
    base_doc = Document(page_content=segment.page_content, metadata=base)
    try:
        final_chunks = split_document(base_doc, chunking_config)
    except Exception as exc:
        print(f"[警告] chunking strategy={chunking_config.strategy} 失败，回退旧 header-aware 切分：{exc}", file=sys.stderr)
        section_docs = header_splitter.split_text(segment.page_content)
        final_chunks = []
        for sec in section_docs:
            merged_meta = {**base, **(sec.metadata or {})}
            merged_meta = _clean_metadata(merged_meta)
            section_limit = int(getattr(recursive_splitter, "_chunk_size", 800))
            if len(sec.page_content) <= section_limit:
                final_chunks.append(Document(page_content=sec.page_content, metadata=merged_meta))
                continue
            parent = Document(page_content=sec.page_content, metadata=merged_meta)
            final_chunks.extend(recursive_splitter.split_documents([parent]))
    out: List[Document] = []
    for c in final_chunks:
        md = _clean_metadata(dict(c.metadata))
        md.setdefault("chunk_strategy", chunking_config.strategy)
        md.setdefault("chunk_index", len(out))
        md.setdefault("section_title", str(md.get("Header 3") or md.get("Header 2") or md.get("Header 1") or "Untitled"))
        md.setdefault("section_type", "other")
        body = c.page_content or ""
        if md.get("doc_type") == "paper":
            body = _paper_embedding_prefix_line(md) + body
        quality = assess_text_quality(body)
        md["text_quality_score"] = quality["score"]
        if quality["warning"]:
            md["text_quality_warning"] = quality["warning"]
            md["text_quality_long_alpha_tokens"] = quality["long_alpha_tokens"]
            md["text_quality_whitespace_ratio"] = quality["whitespace_ratio"]
        out.append(Document(page_content=body, metadata=md))
    return out


def build_embeddings():
    provider = (os.getenv("EMBEDDING_PROVIDER") or "google").lower().strip()

    if provider in ("local_hash", "hash", "offline"):
        return LocalHashEmbeddings(int(os.getenv("LOCAL_HASH_EMBEDDING_DIM", "384")))

    if provider in ("google", "gemini"):
        if GoogleGenerativeAIEmbeddings is None:
            raise ValueError("未安装 langchain-google-genai，无法使用 Google 嵌入。")
        api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("使用 Google 嵌入时请设置 GOOGLE_API_KEY 或 GEMINI_API_KEY")
        model = os.getenv("GOOGLE_EMBEDDING_MODEL", "gemini-embedding-001")
        # 不在构造函数里传 task_type：LangChain 对 embed_documents 默认 RETRIEVAL_DOCUMENT（入库），
        # 对 embed_query 默认 RETRIEVAL_QUERY（检索）。若此处写死 task_type=RETRIEVAL_DOCUMENT，
        # 检索时查询向量也会变成 DOCUMENT 任务类型，损害相似度。
        return GoogleGenerativeAIEmbeddings(model=model, google_api_key=api_key)

    if provider == "zhipu":
        api_key = os.getenv("ZHIPU_API_KEY")
        if not api_key:
            raise ValueError("EMBEDDING_PROVIDER=zhipu 时必须设置 ZHIPU_API_KEY")
        return OpenAIEmbeddings(
            model=os.getenv("ZHIPU_EMBEDDING_MODEL", "embedding-2"),
            openai_api_key=api_key,
            openai_api_base=os.getenv("ZHIPU_OPENAI_BASE", "https://open.bigmodel.cn/api/paas/v4/"),
        )

    if provider in ("huggingface", "bge_m3", "e5"):
        if HuggingFaceEmbeddings is None:
            raise ValueError("未安装 sentence-transformers / langchain-community，无法使用 huggingface 嵌入。")
        if provider == "bge_m3":
            model_name = os.getenv("HF_EMBEDDING_MODEL") or os.getenv("BGE_M3_EMBEDDING_MODEL", "BAAI/bge-m3")
        elif provider == "e5":
            model_name = os.getenv("E5_EMBEDDING_MODEL", "intfloat/multilingual-e5-base")
        else:
            model_name = os.getenv("HF_EMBEDDING_MODEL", "BAAI/bge-large-zh-v1.5")
        normalize = (os.getenv("EMBEDDING_NORMALIZE", "true") or "").lower() not in ("0", "false", "no")
        return HuggingFaceEmbeddings(model_name=model_name, encode_kwargs={"normalize_embeddings": normalize})

    if provider == "openai":
        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("EMBEDDING_PROVIDER=openai 时请设置 OPENAI_API_KEY")
        return OpenAIEmbeddings(model=os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small"))

    raise ValueError(f"不支持的 EMBEDDING_PROVIDER: {provider}")


# ---------- 增量更新：processed_files.json + Chroma 按 source 去重 ----------


def compute_file_md5(path: Path, chunk_bytes: int = 1024 * 1024) -> str:
    """对文件二进制内容计算 MD5，用于判断文件是否被替换或编辑。"""
    h = hashlib.md5()
    with path.open("rb") as f:
        while True:
            block = f.read(chunk_bytes)
            if not block:
                break
            h.update(block)
    return h.hexdigest()


def load_processed_registry() -> Dict[str, Dict[str, Any]]:
    """
    读取处理记录。结构示例：
    {"version": 1, "files": {"papers/论文.pdf": {"md5": "...", "mtime": 1710000000.0}}}
    """
    if not PROCESSED_FILES_JSON.is_file():
        return {"version": 1, "files": {}}
    try:
        data = json.loads(PROCESSED_FILES_JSON.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            return {"version": 1, "files": {}}
        files = data.get("files")
        if not isinstance(files, dict):
            files = {}
        return {"version": int(data.get("version", 1)), "files": dict(files)}
    except (json.JSONDecodeError, OSError, ValueError):
        print("[警告] processed_files.json 损坏或不可读，将视为空记录。", file=sys.stderr)
        return {"version": 1, "files": {}}


def save_processed_registry(registry: Dict[str, Any]) -> None:
    """原子写入处理记录，避免中断导致 JSON 半截。"""
    registry.setdefault("version", 1)
    registry.setdefault("files", {})
    tmp = PROCESSED_FILES_JSON.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(registry, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(PROCESSED_FILES_JSON)


def load_corpus_manifest() -> Dict[str, Any]:
    """Read corpus manifest; created/updated during ingest for reproducibility."""
    if not CORPUS_MANIFEST_JSON.is_file():
        return {"version": 1, "files": {}}
    try:
        data = json.loads(CORPUS_MANIFEST_JSON.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            return {"version": 1, "files": {}}
        files = data.get("files")
        if not isinstance(files, dict):
            files = {}
        return {"version": int(data.get("version", 1)), "files": dict(files)}
    except (json.JSONDecodeError, OSError, ValueError):
        print("[警告] corpus_manifest.json 损坏或不可读，将重新生成。", file=sys.stderr)
        return {"version": 1, "files": {}}


def save_corpus_manifest(manifest: Dict[str, Any]) -> None:
    manifest.setdefault("version", 1)
    manifest.setdefault("files", {})
    tmp = CORPUS_MANIFEST_JSON.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(CORPUS_MANIFEST_JSON)


def classify_pdf_for_ingest(
    file_path: Path,
    registry: Dict[str, Any],
) -> Literal["new", "updated", "skipped"]:
    """
    单文件增量判定（键为相对 data/ 的路径，与 metadata['source'] 一致）：

    - 记录不存在 → new
    - md5 与 mtime 均与记录一致 → skipped（不重复调用解析/嵌入 API）
    - 否则（内容变或仅 touch 导致 mtime 变）→ updated，需先删向量再重灌
    """
    key = pdf_source_key(file_path)
    files: Dict[str, Any] = registry.get("files") or {}
    rec = files.get(key)
    md5_now = compute_file_md5(file_path)
    mtime_now = float(file_path.stat().st_mtime)
    if not rec:
        return "new"
    try:
        md5_old = str(rec.get("md5", ""))
        mtime_old = float(rec.get("mtime", -1.0))
    except (TypeError, ValueError):
        return "updated"
    # JSON 读写后 mtime 可能有浮点误差，使用小容差
    if md5_old == md5_now and abs(mtime_old - mtime_now) < 1e-6:
        return "skipped"
    return "updated"


def record_pdf_processed(registry: Dict[str, Any], file_path: Path) -> None:
    """在内存中标记某文件已成功入库（调用方负责随后 save_processed_registry）。"""
    files = registry.setdefault("files", {})
    files[pdf_source_key(file_path)] = {
        "md5": compute_file_md5(file_path),
        "mtime": float(file_path.stat().st_mtime),
    }


def _parser_hint_for_file(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return "llamaparse_or_python_docx" if os.getenv("LLAMA_CLOUD_API_KEY") else "python_docx"
    if os.getenv("LLAMA_CLOUD_API_KEY"):
        return "llamaparse"
    if (os.getenv("INGEST_PDFPLUMBER_FALLBACK", "") or "").lower() in ("1", "true", "yes"):
        return "pdfplumber"
    return "unconfigured"


def _quality_summary_for_chunks(chunks: Sequence[Document]) -> Dict[str, Any]:
    warnings = Counter()
    scores: List[float] = []
    for doc in chunks:
        meta = doc.metadata or {}
        try:
            scores.append(float(meta.get("text_quality_score", 1.0)))
        except (TypeError, ValueError):
            pass
        warning = str(meta.get("text_quality_warning") or "").strip()
        if warning:
            for part in warning.split(","):
                if part:
                    warnings[part] += 1
    low_quality = sum(1 for s in scores if s < 0.75)
    avg_score = round(sum(scores) / len(scores), 3) if scores else None
    return {
        "chunk_count": len(chunks),
        "avg_text_quality_score": avg_score,
        "low_quality_chunk_count": low_quality,
        "warnings": dict(warnings),
    }


def record_corpus_manifest_entry(
    manifest: Dict[str, Any],
    file_path: Path,
    *,
    doc_type: str,
    doc_role: str,
    chunks: Sequence[Document],
    global_metadata: Dict[str, Any],
    pairing: Optional[Dict[str, Any]] = None,
) -> None:
    """Record reproducibility metadata for a successfully ingested source."""
    key = pdf_source_key(file_path)
    files = manifest.setdefault("files", {})
    files[key] = {
        "source": key,
        "filename": file_path.name,
        "suffix": file_path.suffix.lower(),
        "doc_type": doc_type,
        "doc_role": doc_role,
        "md5": compute_file_md5(file_path),
        "mtime": float(file_path.stat().st_mtime),
        "parser_hint": _parser_hint_for_file(file_path),
        "chunking": {
            "strategy": "markdown_headers_then_recursive",
            "paper_chunk_size": PAPER_CHUNK_SIZE if doc_type == "paper" else None,
            "paper_chunk_overlap": PAPER_CHUNK_OVERLAP if doc_type == "paper" else None,
            "sop_chunk_size": SOP_CHUNK_SIZE if doc_type == "sop" else None,
            "sop_chunk_overlap": SOP_CHUNK_OVERLAP if doc_type == "sop" else None,
        },
        "embedding": {
            "provider": (os.getenv("EMBEDDING_PROVIDER") or "google").lower().strip(),
            "model": (
                os.getenv("GOOGLE_EMBEDDING_MODEL")
                or os.getenv("OPENAI_EMBEDDING_MODEL")
                or os.getenv("ZHIPU_EMBEDDING_MODEL")
                or os.getenv("HF_EMBEDDING_MODEL")
                or "gemini-embedding-001"
            ),
        },
        "metadata": {
            "project_id": global_metadata.get("project_id"),
            "paper_title": global_metadata.get("paper_title"),
            "publication_year": global_metadata.get("publication_year"),
        },
        "pairing": pairing or {},
        "quality": _quality_summary_for_chunks(chunks),
    }


def prune_corpus_manifest_for_removed_files(disk_paths: Sequence[Path], manifest: Dict[str, Any]) -> int:
    names_on_disk = {pdf_source_key(p) for p in disk_paths}
    files: Dict[str, Any] = dict(manifest.get("files") or {})
    stale = [name for name in list(files.keys()) if name not in names_on_disk]
    for name in stale:
        files.pop(name, None)
    manifest["files"] = files
    return len(stale)


def backfill_corpus_manifest_from_chroma(
    manifest: Dict[str, Any],
    disk_paths: Sequence[Path],
    vectorstore: Chroma,
) -> int:
    """
    Populate manifest entries from existing Chroma metadata when no re-ingest is needed.

    New quality metadata is only available after re-ingest, but this still records
    source/doc_type/chunk counts/model config for the current persisted index.
    """
    path_by_source = {pdf_source_key(p): p for p in disk_paths}
    files = manifest.setdefault("files", {})
    missing_keys = [k for k in path_by_source if k not in files]
    if not missing_keys:
        return 0
    try:
        col = vectorstore._collection  # noqa: SLF001 - manifest backfill from persisted metadata
        batch = col.get(include=["metadatas"], limit=200_000)
    except Exception as exc:
        print(f"[警告] 无法从 Chroma 回填 corpus_manifest：{exc}", file=sys.stderr)
        return 0
    grouped: Dict[str, List[Document]] = {}
    for meta in batch.get("metadatas") or []:
        if not isinstance(meta, dict):
            continue
        src = str(meta.get("source") or "").strip()
        if src in path_by_source:
            grouped.setdefault(src, []).append(Document(page_content="", metadata=meta))
    added = 0
    for src in missing_keys:
        docs = grouped.get(src) or []
        if not docs:
            continue
        path = path_by_source[src]
        first_meta = docs[0].metadata or {}
        doc_type = str(first_meta.get("doc_type") or ("sop" if src.startswith("manuals/") else "paper"))
        doc_role = str(first_meta.get("doc_role") or ("manual" if doc_type == "sop" else "unknown"))
        pairing = {
            "paired_main_source": first_meta.get("paired_main_source"),
            "paired_supplementary_sources": first_meta.get("paired_supplementary_sources"),
            "backfilled_from_chroma": True,
        }
        record_corpus_manifest_entry(
            manifest,
            path,
            doc_type=doc_type,
            doc_role=doc_role,
            chunks=docs,
            global_metadata=first_meta,
            pairing=pairing,
        )
        added += 1
    return added


def open_chroma_store(embeddings) -> Chroma:
    """打开（或创建）持久化 Chroma 集合，用于增量 add_documents / delete。"""
    CHROMA_PERSIST_DIR.mkdir(parents=True, exist_ok=True)
    return Chroma(
        persist_directory=str(CHROMA_PERSIST_DIR),
        embedding_function=embeddings,
        collection_name=CHROMA_COLLECTION_NAME,
    )


def delete_chunks_by_source(vectorstore: Chroma, source_key: str) -> None:
    """
    按 metadata.source 精确删除某文件对应的全部向量行。

    说明：source_key 为相对 data/ 的路径（如 papers/a.pdf、manuals/SOP.docx），与入库时 metadata['source'] 一致。
    """
    if not source_key:
        return
    try:
        vectorstore.delete(where={"source": source_key})
    except Exception as exc:  # pragma: no cover
        print(f"[警告] 按 source 删除向量失败：{source_key} — {exc}", file=sys.stderr)


def prune_registry_and_chroma_for_removed_files(
    disk_paths: Sequence[Path],
    registry: Dict[str, Any],
    vectorstore: Chroma,
) -> int:
    """
    磁盘上已不存在的文件，若仍留在处理记录中，则同步删除 Chroma 中对应 source 并清理记录。
    返回被清理的键数量。
    """
    names_on_disk = {pdf_source_key(p) for p in disk_paths}
    files: Dict[str, Any] = dict(registry.get("files") or {})
    stale = [name for name in list(files.keys()) if name not in names_on_disk]
    for name in stale:
        delete_chunks_by_source(vectorstore, name)
        files.pop(name, None)
        print(f"  [清理] 磁盘已删除文件，已移除向量与记录：{name}")
    registry["files"] = files
    return len(stale)


def _process_single_pdf_to_chunks(
    file_path: Path,
    header_splitter: MarkdownHeaderTextSplitter,
    recursive_splitter: RecursiveCharacterTextSplitter,
    global_file_metadata: Dict[str, Any],
    *,
    source_key: str,
    chunking_config: Optional[ChunkingConfig] = None,
) -> List[Document]:
    segments = load_document_markdown_segments(file_path, source_key)
    chunks: List[Document] = []
    for seg in segments:
        chunks.extend(
            hierarchical_chunk_markdown_segment(
                seg,
                header_splitter,
                recursive_splitter,
                global_file_metadata=global_file_metadata,
                chunking_config=chunking_config,
            )
        )
    return chunks


def ingest_all_pdfs() -> int:
    """
    分流入库：papers/ → doc_type=paper；manuals/ → doc_type=sop。
    论文走 Main+SI 单元与 LLM 题录；手册走大 chunk、无 SI 配对。
    增量规则见 `processed_files.json`；更新时按 metadata.source（相对路径）删除旧向量。
    """
    ensure_ingest_directories()

    jobs = iter_ingest_jobs()
    all_paths = [p for p, _ in jobs]

    if DATA_DIR.is_dir():
        for p in DATA_DIR.iterdir():
            if p.is_file() and is_supported_ingest_file(p):
                print(
                    "[提示] 发现位于 data/ 根目录的文档，请移入 data/papers/（论文）或 data/manuals/（操作手册）后再运行 ingest。",
                    file=sys.stderr,
                )
                break

    if not all_paths:
        print(f"在 {DATA_PAPERS_DIR} 与 {DATA_MANUALS_DIR} 中未找到支持的文档（{', '.join(sorted(INGEST_SUPPORTED_SUFFIXES))}）。")
        return 0

    registry = load_processed_registry()
    manifest = load_corpus_manifest()
    rebuild = (os.getenv("REBUILD_CHROMA", "") or "").lower() in ("1", "true", "yes")
    if rebuild:
        if CHROMA_PERSIST_DIR.exists():
            shutil.rmtree(CHROMA_PERSIST_DIR)
            print(f"已按 REBUILD_CHROMA 清空向量库目录：{CHROMA_PERSIST_DIR}")
        registry = {"version": 1, "files": {}}
        manifest = {"version": 1, "files": {}}
        print("已重置 processed_files 记录（与全量重建对齐）。")
        save_processed_registry(registry)
        save_corpus_manifest(manifest)

    embeddings = build_embeddings()
    vectorstore = open_chroma_store(embeddings)

    prune_registry_and_chroma_for_removed_files(all_paths, registry, vectorstore)
    pruned_manifest = prune_corpus_manifest_for_removed_files(all_paths, manifest)
    if pruned_manifest:
        print(f"  [清理] 已从 corpus_manifest 移除 {pruned_manifest} 个不存在的文件记录")
    backfilled_manifest = backfill_corpus_manifest_from_chroma(manifest, all_paths, vectorstore)
    if backfilled_manifest:
        print(f"  [manifest] 已从现有 Chroma metadata 回填 {backfilled_manifest} 个文件记录")
    save_processed_registry(registry)
    save_corpus_manifest(manifest)

    ingest_status: Dict[str, Literal["new", "updated", "skipped"]] = {
        pdf_source_key(p): classify_pdf_for_ingest(p, registry) for p in all_paths
    }
    new_n = sum(1 for s in ingest_status.values() if s == "new")
    upd_n = sum(1 for s in ingest_status.values() if s == "updated")
    skip_n = sum(1 for s in ingest_status.values() if s == "skipped")
    print(f"发现 {new_n} 个新文档，{upd_n} 个更新文档，跳过 {skip_n} 个已存在文档")

    if new_n + upd_n == 0:
        save_processed_registry(registry)
        save_corpus_manifest(manifest)
        print("无需写入向量库（全部文件均已是最新指纹）。")
        return 0

    header_splitter = _build_markdown_header_splitter()
    paper_recursive = _build_recursive_splitter(PAPER_CHUNK_SIZE, PAPER_CHUNK_OVERLAP)
    sop_recursive = _build_recursive_splitter(SOP_CHUNK_SIZE, SOP_CHUNK_OVERLAP)
    paper_chunking = ChunkingConfig(
        strategy=PAPER_CHUNK_STRATEGY,
        chunk_size=PAPER_CHUNK_SIZE,
        chunk_overlap=PAPER_CHUNK_OVERLAP,
    )
    sop_chunking = ChunkingConfig(
        strategy=SOP_CHUNK_STRATEGY,
        chunk_size=SOP_CHILD_CHUNK_SIZE,
        chunk_overlap=SOP_CHILD_CHUNK_OVERLAP,
        child_chunk_size=SOP_CHILD_CHUNK_SIZE,
        child_chunk_overlap=SOP_CHILD_CHUNK_OVERLAP,
        parent_chunk_size=SOP_PARENT_CHUNK_SIZE,
        parent_chunk_overlap=SOP_PARENT_CHUNK_OVERLAP,
    )
    print(f"Chunking: paper={paper_chunking.strategy} ({PAPER_CHUNK_SIZE}/{PAPER_CHUNK_OVERLAP}), sop={sop_chunking.strategy}")
    total_new_chunks = 0

    # ---------- 论文（papers/）：单元 + SI ----------
    paper_paths = [p for p, t in jobs if t == "paper"]
    units = group_pdfs_into_paper_units(paper_paths)
    warn_paper_unit_pairing_anomalies(units)

    for unit in units:
        files_in_unit: List[Path] = []
        if unit.main_pdf is not None:
            files_in_unit.append(unit.main_pdf)
        files_in_unit.extend(unit.extra_main_pdfs)
        files_in_unit.extend(unit.supplementary_pdfs)

        files_work = [p for p in files_in_unit if ingest_status.get(pdf_source_key(p)) != "skipped"]
        if not files_work:
            continue

        label = unit.main_pdf.name if unit.main_pdf else (unit.supplementary_pdfs[0].name if unit.supplementary_pdfs else unit.core_stem)
        print(f"[论文] 单元 [{unit.core_stem}] — 待处理: {len(files_work)} 个文件 — {label}")

        for p in files_work:
            if ingest_status.get(pdf_source_key(p)) == "updated":
                sk = pdf_source_key(p)
                delete_chunks_by_source(vectorstore, sk)
                print(f"  [更新] 已删除旧向量：source={sk}")

        anchor = unit.main_pdf if unit.main_pdf is not None else files_work[0]
        anchor_key = pdf_source_key(anchor)
        try:
            anchor_segments = load_document_markdown_segments(anchor, anchor_key)
        except Exception as e:
            print(f"  [失败] 无法解析锚点文件 {anchor_key}: {e}", file=sys.stderr)
            continue

        try:
            llm_meta, project_id = extract_paper_metadata_llm(anchor_segments, anchor)
        except Exception as e:
            print(f"  [警告] 元数据流程异常，使用回退：{e}", file=sys.stderr)
            llm_meta, project_id = _fallback_metadata_from_path(anchor)

        flat_llm = _flatten_paper_meta_for_chroma(llm_meta, project_id)

        if unit.main_pdf is not None and unit.main_pdf in files_work:
            try:
                sup_keys = " | ".join(pdf_source_key(p) for p in unit.supplementary_pdfs)
                meta_main = {
                    **flat_llm,
                    "doc_type": "paper",
                    "doc_role": "main_text",
                    "paired_supplementary_sources": (sup_keys[:2000] if sup_keys else ""),
                }
                mk = pdf_source_key(unit.main_pdf)
                ch = _process_single_pdf_to_chunks(
                    unit.main_pdf,
                    header_splitter,
                    paper_recursive,
                    meta_main,
                    source_key=mk,
                    chunking_config=paper_chunking,
                )
                if ch:
                    vectorstore.add_documents(ch)
                    total_new_chunks += len(ch)
                record_pdf_processed(registry, unit.main_pdf)
                record_corpus_manifest_entry(
                    manifest,
                    unit.main_pdf,
                    doc_type="paper",
                    doc_role="main_text",
                    chunks=ch,
                    global_metadata=meta_main,
                    pairing={
                        "core_stem": unit.core_stem,
                        "paired_supplementary_sources": [pdf_source_key(p) for p in unit.supplementary_pdfs],
                        "extra_main_sources": [pdf_source_key(p) for p in unit.extra_main_pdfs],
                    },
                )
                save_processed_registry(registry)
                save_corpus_manifest(manifest)
                print(f"  正文 {mk}: 写入 {len(ch)} chunks")
            except Exception as e:
                print(f"  [失败] 正文切分失败 {unit.main_pdf.name}: {e}", file=sys.stderr)

        for extra_main in unit.extra_main_pdfs:
            if extra_main not in files_work:
                continue
            try:
                sup_keys = " | ".join(pdf_source_key(p) for p in unit.supplementary_pdfs)
                meta_extra_main = {
                    **flat_llm,
                    "doc_type": "paper",
                    "doc_role": "main_text",
                    "paired_supplementary_sources": (sup_keys[:2000] if sup_keys else ""),
                    "duplicate_main_anchor_source": pdf_source_key(unit.main_pdf) if unit.main_pdf else "",
                }
                ek = pdf_source_key(extra_main)
                ch_extra = _process_single_pdf_to_chunks(
                    extra_main,
                    header_splitter,
                    paper_recursive,
                    meta_extra_main,
                    source_key=ek,
                    chunking_config=paper_chunking,
                )
                if ch_extra:
                    vectorstore.add_documents(ch_extra)
                    total_new_chunks += len(ch_extra)
                record_pdf_processed(registry, extra_main)
                record_corpus_manifest_entry(
                    manifest,
                    extra_main,
                    doc_type="paper",
                    doc_role="main_text",
                    chunks=ch_extra,
                    global_metadata=meta_extra_main,
                    pairing={
                        "core_stem": unit.core_stem,
                        "duplicate_main_anchor_source": pdf_source_key(unit.main_pdf) if unit.main_pdf else "",
                        "paired_supplementary_sources": [pdf_source_key(p) for p in unit.supplementary_pdfs],
                    },
                )
                save_processed_registry(registry)
                save_corpus_manifest(manifest)
                print(f"  额外正文 {ek}: 写入 {len(ch_extra)} chunks")
            except Exception as e:
                print(f"  [失败] 额外正文切分失败 {extra_main.name}: {e}", file=sys.stderr)

        for si_path in unit.supplementary_pdfs:
            if si_path not in files_work:
                continue
            try:
                main_key = pdf_source_key(unit.main_pdf) if unit.main_pdf else ""
                meta_si = {
                    **flat_llm,
                    "doc_type": "paper",
                    "doc_role": "supplementary_info",
                    "paired_main_source": main_key,
                }
                sk = pdf_source_key(si_path)
                ch_si = _process_single_pdf_to_chunks(
                    si_path, header_splitter, paper_recursive, meta_si, source_key=sk, chunking_config=paper_chunking
                )
                if ch_si:
                    vectorstore.add_documents(ch_si)
                    total_new_chunks += len(ch_si)
                record_pdf_processed(registry, si_path)
                record_corpus_manifest_entry(
                    manifest,
                    si_path,
                    doc_type="paper",
                    doc_role="supplementary_info",
                    chunks=ch_si,
                    global_metadata=meta_si,
                    pairing={"core_stem": unit.core_stem, "paired_main_source": main_key},
                )
                save_processed_registry(registry)
                save_corpus_manifest(manifest)
                print(f"  补充 {sk}: 写入 {len(ch_si)} chunks")
            except Exception as e:
                print(f"  [失败] 补充材料切分失败 {si_path.name}: {e}", file=sys.stderr)

    # ---------- 操作手册（manuals/）：大 chunk + doc_type=sop ----------
    for path, dtype in jobs:
        if dtype != "sop":
            continue
        sk = pdf_source_key(path)
        if ingest_status.get(sk) == "skipped":
            continue
        print(f"[手册] 处理 {sk}")
        if ingest_status.get(sk) == "updated":
            delete_chunks_by_source(vectorstore, sk)
            print(f"  [更新] 已删除旧向量：source={sk}")
        try:
            meta_sop = build_sop_global_metadata(path)
            ch = _process_single_pdf_to_chunks(
                path, header_splitter, sop_recursive, meta_sop, source_key=sk, chunking_config=sop_chunking
            )
            if ch:
                vectorstore.add_documents(ch)
                total_new_chunks += len(ch)
            record_pdf_processed(registry, path)
            record_corpus_manifest_entry(
                manifest,
                path,
                doc_type="sop",
                doc_role="manual",
                chunks=ch,
                global_metadata=meta_sop,
                pairing={},
            )
            save_processed_registry(registry)
            save_corpus_manifest(manifest)
            print(f"  手册 {sk}: 写入 {len(ch)} chunks（chunk={SOP_CHUNK_SIZE}/{SOP_CHUNK_OVERLAP}）")
        except Exception as e:
            print(f"  [失败] 手册切分失败 {sk}: {e}", file=sys.stderr)

    print(
        f"完成。本次新增/更新向量片段数：{total_new_chunks}；"
        f"向量库路径：{CHROMA_PERSIST_DIR}；集合名：{CHROMA_COLLECTION_NAME}"
    )
    return total_new_chunks


def main() -> None:
    try:
        ensure_ingest_directories()
        jobs_existing = iter_ingest_jobs()
        n = ingest_all_pdfs()
        if len(jobs_existing) == 0 and n == 0:
            sys.exit(1)
        sys.exit(0)
    except Exception as e:
        print(f"ingest 失败：{e}", file=sys.stderr)
        sys.exit(2)


if __name__ == "__main__":
    main()
